from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

def set_run(run, bold=False, size=11, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(0)
    return p

# ── Title block
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('C-ECM \u2014 Centralized Enterprise Content Management')
set_run(r, bold=True, size=20, color=(31, 35, 40))

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('IBM TechXchange 2026 Pre-conference Dev Day Hackathon  \u00b7  #watsonxHackathon')
set_run(r2, size=11, color=(87, 96, 106))

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = meta.add_run('3-Minute Video Script  \u00b7  ~420 words  \u00b7  ~130 words/min')
set_run(r3, size=10, color=(87, 96, 106))

doc.add_paragraph()
add_rule(doc)
doc.add_paragraph()

beats = [
    {
        'time':  '0:00 \u2013 0:20',
        'title': 'BEAT 0 \u2014 Welcome',
        'color': (15, 118, 110),
        'lines': [
            'Hi everyone! Welcome to the IBM Bob Hackathon \u2014 I\u2019m Mohammad Jamil Ahmed, and I\u2019m incredibly excited to be here.',
            'I want to be honest with you. During this build, IBM Bob hit its 40,000-token context limit \u2014 multiple times. The session would fill up, Bob would lose track of earlier decisions, and I had to export the task, start a fresh session, and re-anchor it with context from the previous one.',
            'But here\u2019s what\u2019s remarkable: even with that constraint, I managed the entire project across five Bob sessions \u2014 exporting tasks, re-loading context, and keeping the build moving forward. That discipline of working with an AI agent across its limits is itself a real-world skill. And the result? Nine-hundred-and-eighty-one Bob interactions, five exported session files, and a fully working product. Let me show you.',
        ]
    },
    {
        'time':  '0:20 \u2013 0:50',
        'title': 'BEAT 1 \u2014 The Problem',
        'color': (59, 130, 212),
        'lines': [
            'Picture a Friday afternoon production incident. Something broke in the last deployment. The on-call engineer needs the approved runbook, the change-approval record, and the audit log \u2014 right now.',
            'What actually happens? The official document is in FileNet \u2014 but nobody knows which version was signed off. The working copy is on a shared drive, in three folders, with three different names. The approval lives in an email thread from six months ago. And the mainframe DMS has never been connected to anything.',
            "That's five systems that don't talk to each other \u2014 and the clock is ticking.  This is not an edge case. This is every bank, insurer, and government agency running enterprise workloads.",
        ]
    },
    {
        'time':  '0:50 \u2013 1:20',
        'title': 'BEAT 2 \u2014 The Solution',
        'color': (22, 163, 74),
        'lines': [
            'I built C-ECM \u2014 Centralized Enterprise Content Management. It\u2019s a governance and intelligence layer that sits in front of every storage system your teams already use \u2014 IBM FileNet, IBM i, IBM Z mainframe, SharePoint, Google Drive, S3, Azure Blob, Box \u2014 eleven backends in total.',
            'It looks and feels like a shared drive. But behind the scenes, every document is now approval-gated, version-controlled, audit-logged, and cross-system-searchable \u2014 on day one, with zero migration required.',
        ]
    },
    {
        'time':  '1:20 \u2013 2:00',
        'title': 'BEAT 3 \u2014 Key Features Demo',
        'color': (124, 92, 216),
        'lines': [
            '[SHOW UI]  One search bar queries all eleven backends in parallel \u2014 not sequentially, so you get results in the time it takes to hit the slowest single system.',
            '[SHOW WORKFLOW]  Multi-step, quorum-based approval workflows \u2014 the same gate a CAB or change-management process needs. Every vote, every rejection, every sign-off is permanently recorded.',
            '[SHOW AUDIT]  The compliance dashboard lets you reconstruct any change-approval trail in under five minutes \u2014 not five days. One-click CSV export for auditors.',
            '[SHOW AI]  And IBM watsonx.ai is integrated for document intelligence \u2014 summarize, classify, and ask natural-language questions about any document without even opening it.',
        ]
    },
    {
        'time':  '2:00 \u2013 2:40',
        'title': 'BEAT 4 \u2014 Built with IBM Bob 2.0',
        'color': (217, 119, 6),
        'lines': [
            'This entire project was built end-to-end using IBM Bob 2.0 in Agent mode \u2014 not as an autocomplete tool, but as a true agentic developer.',
            'Bob read all three official hackathon PDFs directly, extracted the judging criteria, and used them to shape the product. It fanned out eight parallel subagents to review the codebase simultaneously. When Global Search broke, Bob traced the bug across three independent layers in one session \u2014 a frontend issue, a backend schema mismatch, and a race condition in the storage registry \u2014 reproduced it deterministically, fixed it, and re-ran a two-hundred-trial stress test to prove the fix.',
            'Every change was verified against a live 217-test backend suite before being reported complete. Nine-hundred-and-eighty-one total IBM Bob interactions are exported and committed to the repo as evidence.',
        ]
    },
    {
        'time':  '2:40 \u2013 3:00',
        'title': 'BEAT 5 \u2014 Close',
        'color': (220, 38, 38),
        'lines': [
            'C-ECM removes the document-governance tax that every enterprise pays every day \u2014 without asking teams to migrate a single file.',
            'One governance layer. Eleven storage systems. Zero migration. Built with IBM Bob 2.0.',
            '[SMILE + PAUSE]  Thank you.',
        ]
    },
]

for beat in beats:
    # Beat header line
    p = doc.add_paragraph()
    tb = p.add_run(beat['time'] + '   ')
    set_run(tb, bold=True, size=10, color=(87, 96, 106))
    tt = p.add_run(beat['title'])
    set_run(tt, bold=True, size=13, color=beat['color'])
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)

    for line in beat['lines']:
        p2 = doc.add_paragraph(style='Normal')
        p2.paragraph_format.left_indent = Inches(0.25)
        p2.paragraph_format.space_after = Pt(6)

        if line.startswith('['):
            bracket_end = line.index(']') + 1
            cue_text  = line[:bracket_end]
            rest_text = line[bracket_end:].strip()
            rc = p2.add_run(cue_text + '  ')
            set_run(rc, bold=True, size=11, color=beat['color'])
            rb = p2.add_run(rest_text)
            set_run(rb, size=11)
        else:
            rb = p2.add_run(line)
            set_run(rb, size=11)

    add_rule(doc)
    doc.add_paragraph()

# Delivery Tips
h = doc.add_heading('Delivery Tips', level=2)

tips = [
    ('Pace',     'Speak slightly slower than feels natural \u2014 ~130 words/min lands well on camera.'),
    ('Beat 1',   'Lean in and sound frustrated. The pain is real \u2014 let it show.'),
    ('Beat 3',   'Have the app running on screen; switch to it on each screen cue.'),
    ('Beat 4',   'Slow down on "981 interactions" \u2014 big number, let it land.'),
    ('Beat 5',   'Say the three-line close slowly. It is your tagline. Own it.'),
    ('If short', 'Cut the last sentence of Beat 3 (the watsonx line) to save ~10 seconds.'),
]
for label, tip in tips:
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(4)
    rl = p.add_run(label + ':  ')
    set_run(rl, bold=True, size=11)
    rt = p.add_run(tip)
    set_run(rt, size=11)

doc.add_paragraph()

# Footer
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = fp.add_run('Mohammad Jamil Ahmed  \u00b7  IBM TechXchange 2026 Hackathon  \u00b7  Built with IBM Bob 2.0')
set_run(rf, size=9, color=(87, 96, 106))

out_path = 'C-ECM_Video_Script_v2.docx'
doc.save(out_path)
print('Saved:', out_path)
